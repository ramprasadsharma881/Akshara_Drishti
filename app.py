"""
AksharaDrishti OCR API - Cloud Vision Level Accuracy
====================================================
High-performance OCR for old Telugu/Kannada Vedantha prints

Features:
- Advanced OpenCV preprocessing (CLAHE, denoising, adaptive threshold)
- Parallel page processing with ProcessPoolExecutor
- Linguistic Repair Layer (Sandhi, Danda normalization)
- Confidence scoring with quality report
- Robust error handling
"""

import os
import time
import uuid
from pathlib import Path
from PIL import Image
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
import shutil
from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing

# ============================================================================
# CONFIGURATION
# ============================================================================

# Number of parallel workers (reduced to prevent thread exhaustion)
# OpenCV + Tesseract each spawn threads, so fewer workers = stable
NUM_WORKERS = min(8, max(2, multiprocessing.cpu_count() // 4))

app = FastAPI(title="AksharaDrishti OCR API - Press Ready")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directories
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Tesseract configuration
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

# Override for Linux deployment
if os.path.exists("/usr/bin/tesseract"):
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

# Poppler configuration
POPPLER_PATH = None
possible_paths = [
    r"C:\poppler\poppler-24.08.0\Library\bin",
    r"C:\Users\Suresh Sharma\OneDrive\Desktop\OCR_Tools\poppler-24.02.0\Library\bin",
    r"C:\Program Files\poppler\Library\bin",
    r"C:\poppler\Library\bin",
]

for path in possible_paths:
    if os.path.exists(path):
        POPPLER_PATH = path
        break

# Store job status
jobs = {}


class OCRJob:
    def __init__(self, job_id: str):
        self.job_id = job_id
        self.status = "processing"
        self.progress = 0
        self.current_page = 0
        self.total_pages = 0
        self.elapsed_time = 0
        self.estimated_time = 0
        self.log = []
        self.output_file = None
        self.error = None
        self.start_time = time.time()
        self.avg_confidence = 0
        self.total_low_confidence = 0
        self.failed_pages = []

    def update(self, **kwargs):
        for key, value in kwargs.items():
            setattr(self, key, value)
        self.elapsed_time = int(time.time() - self.start_time)

    def add_log(self, message: str):
        self.log.append(message)


# Import OCR utilities
from ocr_utils import (
    process_single_page,
    add_paragraphs_to_document,
    add_quality_report,
    preprocess_image
)

# Import AI Spell Check (OpenAI GPT)
from openai_spell_check import spell_check_page_text, SPELL_CHECK_ENABLED


def perform_ocr(job_id: str, pdf_path: str, language: str, dpi: int):
    """
    Perform OCR on PDF using PARALLEL page processing.
    Includes Linguistic Repair Layer and Quality Report.
    """
    job = jobs[job_id]
    
    try:
        job.add_log("═" * 50)
        job.add_log("AksharaDrishti - Press Ready OCR")
        job.add_log("═" * 50)
        job.add_log(f"File: {os.path.basename(pdf_path)}")
        job.add_log(f"Language: {language}")
        job.add_log(f"DPI: {dpi}")
        job.add_log(f"Parallel Workers: {NUM_WORKERS}")
        job.add_log("")
        job.add_log("Features enabled:")
        job.add_log("  • Advanced preprocessing (CLAHE + Denoising)")
        job.add_log("  • Linguistic Repair (Sandhi + Danda normalization)")
        job.add_log("  • Confidence scoring")
        job.add_log("  • Quality report generation")
        job.add_log("")
        
        # Convert PDF to images
        job.add_log("Converting PDF to images...")
        images = convert_from_path(pdf_path, dpi=dpi, poppler_path=POPPLER_PATH)
        job.total_pages = len(images)
        job.update(total_pages=len(images))
        job.add_log(f"Found {len(images)} pages")
        job.add_log("")
        
        # Prepare tasks
        job.add_log(f"Processing {len(images)} pages in parallel...")
        job.add_log("")
        
        tasks = [(i, img, language) for i, img in enumerate(images)]
        
        # Process pages in parallel with error handling
        results = {}
        completed = 0
        total_confidence = 0
        total_low_confidence = 0
        failed_pages = []
        
        with ProcessPoolExecutor(max_workers=NUM_WORKERS) as executor:
            future_to_page = {
                executor.submit(process_single_page, task): task[0] 
                for task in tasks
            }
            
            for future in as_completed(future_to_page):
                page_idx = future_to_page[future]
                
                try:
                    result = future.result(timeout=300)  # 5 min timeout per page
                    results[page_idx] = result
                    
                    completed += 1
                    
                    if result['success']:
                        if result['confidence'] > 0:
                            total_confidence += result['confidence']
                        total_low_confidence += result.get('low_confidence_count', 0)
                        
                        conf = f"{result['confidence']:.0f}%" if result['confidence'] > 0 else "N/A"
                        low_conf = result.get('low_confidence_count', 0)
                        job.add_log(f"  Page {page_idx + 1}: OK (conf: {conf}, uncertain: {low_conf})")
                    else:
                        error_msg = result.get('error', 'Unknown error')
                        failed_pages.append((page_idx, error_msg))
                        job.add_log(f"  Page {page_idx + 1}: FAILED - {error_msg}")
                    
                except Exception as e:
                    # Handle timeout or other errors
                    error_msg = str(e)
                    results[page_idx] = {
                        'page_idx': page_idx,
                        'paragraphs': [],
                        'success': False,
                        'error': error_msg
                    }
                    failed_pages.append((page_idx, error_msg))
                    job.add_log(f"  Page {page_idx + 1}: FAILED - {error_msg}")
                    completed += 1
                
                # Update progress
                job.current_page = completed
                job.progress = (completed / job.total_pages) * 100
                
                elapsed = time.time() - job.start_time
                job.elapsed_time = int(elapsed)
                avg_per_page = elapsed / completed if completed > 0 else 0
                remaining = job.total_pages - completed
                job.estimated_time = int(avg_per_page * remaining)
                job.update()
        
        # Calculate statistics
        successful_pages = len(images) - len(failed_pages)
        job.avg_confidence = total_confidence / successful_pages if successful_pages > 0 else 0
        job.total_low_confidence = total_low_confidence
        job.failed_pages = failed_pages
        
        job.add_log("")
        
        # ====================================================================
        # AI SPELL CHECK (Post-OCR correction using OpenAI GPT)
        # ====================================================================
        # Note: We batch pages together for efficient processing.
        # ====================================================================
        spell_check_applied = 0
        spell_check_skipped = 0
        
        if SPELL_CHECK_ENABLED:
            job.add_log("─" * 40)
            job.add_log("AI Spell Check (OpenAI GPT)")
            job.add_log("─" * 40)
            job.add_log("Processing pages...")
            
            # Batch pages together to reduce API calls
            # Combine every 5 pages into one batch
            PAGES_PER_BATCH = 5
            total_pages = len(images)
            
            for batch_start in range(0, total_pages, PAGES_PER_BATCH):
                batch_end = min(batch_start + PAGES_PER_BATCH, total_pages)
                batch_text_parts = []
                batch_page_indices = []
                
                # Collect text from pages in this batch
                for page_idx in range(batch_start, batch_end):
                    result = results.get(page_idx)
                    if not result or not result['success'] or not result['paragraphs']:
                        continue
                    
                    page_text = '\n\n'.join([p['text'] for p in result['paragraphs']])
                    if len(page_text.strip()) >= 20:
                        batch_text_parts.append(f"[PAGE {page_idx + 1}]\n{page_text}")
                        batch_page_indices.append(page_idx)
                
                if not batch_text_parts:
                    continue
                
                # Combine batch into single text with page markers
                combined_text = '\n\n---PAGE_BREAK---\n\n'.join(batch_text_parts)
                
                job.add_log(f"  Batch {batch_start//PAGES_PER_BATCH + 1}: Pages {batch_start + 1}-{batch_end}...")
                
                try:
                    corrected_text, was_corrected, status = spell_check_page_text(combined_text)
                    
                    if was_corrected:
                        # Split corrected text back to pages
                        corrected_parts = corrected_text.split('---PAGE_BREAK---')
                        
                        for i, page_idx in enumerate(batch_page_indices):
                            if i < len(corrected_parts):
                                result = results.get(page_idx)
                                if result and result['paragraphs']:
                                    # Extract just the text (remove [PAGE X] marker)
                                    corrected_page = corrected_parts[i].strip()
                                    if corrected_page.startswith('[PAGE'):
                                        # Remove the [PAGE X] marker
                                        newline_pos = corrected_page.find('\n')
                                        if newline_pos > 0:
                                            corrected_page = corrected_page[newline_pos:].strip()
                                    
                                    corrected_paragraphs = corrected_page.split('\n\n')
                                    for j, para_data in enumerate(result['paragraphs']):
                                        if j < len(corrected_paragraphs):
                                            para_data['text'] = corrected_paragraphs[j].strip()
                        
                        spell_check_applied += len(batch_page_indices)
                        job.add_log(f"    -> Corrected {len(batch_page_indices)} pages")
                    else:
                        spell_check_skipped += len(batch_page_indices)
                        job.add_log(f"    -> No changes needed")
                        
                except Exception as e:
                    spell_check_skipped += len(batch_page_indices)
                    job.add_log(f"    -> Error: {str(e)[:50]}")
            
            job.add_log(f"")
            job.add_log(f"Spell check: {spell_check_applied} pages corrected, {spell_check_skipped} unchanged")
        else:
            job.add_log("AI Spell Check: DISABLED (set SPELL_CHECK_ENABLED=true to enable)")
        
        job.add_log("")
        job.add_log("Assembling document...")
        
        # Create Word document
        doc = Document()
        
        # Choose font based on language
        if "tel" in language:
            font_name = "Noto Sans Telugu"
        elif "san" in language:
            font_name = "Noto Sans Devanagari"
        else:
            font_name = "Noto Sans Kannada"
        
        # Set up document styles
        normal_style = doc.styles['Normal']
        normal_style.font.name = font_name
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.space_before = Pt(0)
        
        # Assemble document in CORRECT PAGE ORDER
        for i in range(len(images)):
            result = results.get(i, {'paragraphs': [], 'success': False})
            
            if result['success'] and result['paragraphs']:
                add_paragraphs_to_document(doc, result['paragraphs'])
            elif not result['success']:
                # Add placeholder for failed page
                p = doc.add_paragraph()
                p.add_run(f"[Page {i + 1} failed to process: {result.get('error', 'Unknown error')}]")
            
            # Add page break (except for last page - we'll add quality report)
            if i < len(images) - 1:
                doc.add_page_break()
        
        # Add Quality Report at the end
        add_quality_report(
            doc, 
            total_low_confidence, 
            len(images), 
            job.avg_confidence,
            failed_pages
        )
        
        # Save output
        suffix = "_TEL" if "tel" in language else "_KAN"
        filename = os.path.basename(pdf_path).replace(".pdf", f"{suffix}_OCR.docx")
        output_path = OUTPUT_DIR / f"{job_id}_{filename}"
        doc.save(str(output_path))
        
        job.output_file = str(output_path)
        job.status = "completed"
        job.progress = 100
        
        job.add_log("")
        job.add_log("═" * 50)
        job.add_log("OCR COMPLETED - PRESS READY!")
        job.add_log("═" * 50)
        job.add_log(f"Output: {filename}")
        job.add_log(f"Total time: {job.elapsed_time}s")
        job.add_log(f"Avg time per page: {job.elapsed_time / job.total_pages:.1f}s")
        job.add_log(f"Avg confidence: {job.avg_confidence:.1f}%")
        job.add_log(f"Low-confidence words: {total_low_confidence}")
        if failed_pages:
            job.add_log(f"Failed pages: {len(failed_pages)}")
        if SPELL_CHECK_ENABLED:
            job.add_log(f"AI Spell Check: {spell_check_applied} pages corrected")
        job.add_log("")
        job.add_log("Quality Report added at end of document.")
        
    except Exception as e:
        import traceback
        job.status = "failed"
        job.error = str(e)
        job.add_log(f"CRITICAL ERROR: {str(e)}")
        job.add_log(f"Traceback: {traceback.format_exc()}")
    
    finally:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)


@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Serve the main HTML page"""
    html_path = Path("static/index.html")
    if html_path.exists():
        return html_path.read_text()
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>AksharaDrishti - Press Ready OCR</title>
        <meta charset="UTF-8">
    </head>
    <body>
        <h1>AksharaDrishti OCR Service</h1>
        <p>Press-Ready OCR for Telugu/Kannada Vedantic texts</p>
        <p>Features: Linguistic Repair, Confidence Scoring, Quality Reports</p>
        <p>API Documentation: <a href="/docs">/docs</a></p>
    </body>
    </html>
    """


@app.post("/api/upload")
async def upload_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    language: str = Form("kan+san"),
    dpi: int = Form(350)
):
    """Upload PDF and start OCR processing"""
    
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    job_id = str(uuid.uuid4())
    
    upload_path = UPLOAD_DIR / f"{job_id}_{file.filename}"
    with open(upload_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    jobs[job_id] = OCRJob(job_id)
    
    background_tasks.add_task(perform_ocr, job_id, str(upload_path), language, dpi)
    
    features = [
        "Advanced preprocessing (CLAHE + Denoising)",
        "Linguistic Repair (Sandhi + Danda)",
        "Confidence scoring",
        "Quality report"
    ]
    
    if SPELL_CHECK_ENABLED:
        features.append("AI Spell Check (Gemini)")
    
    return {
        "job_id": job_id,
        "message": "OCR processing started (Press Ready mode)",
        "filename": file.filename,
        "workers": NUM_WORKERS,
        "spell_check_enabled": SPELL_CHECK_ENABLED,
        "features": features
    }


@app.get("/api/status/{job_id}")
async def get_status(job_id: str):
    """Get job status and progress"""
    
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    
    return {
        "job_id": job.job_id,
        "status": job.status,
        "progress": job.progress,
        "current_page": job.current_page,
        "total_pages": job.total_pages,
        "elapsed_time": job.elapsed_time,
        "estimated_time": job.estimated_time,
        "avg_confidence": getattr(job, 'avg_confidence', 0),
        "total_low_confidence": getattr(job, 'total_low_confidence', 0),
        "failed_pages": len(getattr(job, 'failed_pages', [])),
        "log": job.log[-20:],
        "error": job.error
    }


@app.get("/api/download/{job_id}")
async def download_result(job_id: str):
    """Download the OCR result"""
    
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    
    if job.status != "completed":
        raise HTTPException(status_code=400, detail="Job not completed yet")
    
    if not job.output_file or not os.path.exists(job.output_file):
        raise HTTPException(status_code=404, detail="Output file not found")
    
    return FileResponse(
        job.output_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(job.output_file)
    )


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    
    # Check if OpenAI API key is configured
    openai_key = os.environ.get("OPENAI_API_KEY")
    
    spell_check_ready = SPELL_CHECK_ENABLED and bool(openai_key)
    
    return {
        "status": "healthy",
        "tesseract": os.path.exists(pytesseract.pytesseract.tesseract_cmd) if pytesseract.pytesseract.tesseract_cmd else False,
        "workers": NUM_WORKERS,
        "features": {
            "preprocessing": "CLAHE + Denoising + Adaptive Threshold",
            "linguistic_repair": "Sandhi + Danda normalization",
            "confidence_scoring": True,
            "quality_report": True,
            "ai_spell_check": spell_check_ready,
            "spell_check_api": "OpenAI GPT" if spell_check_ready else "Missing OPENAI_API_KEY"
        }
    }


# Mount static files
if Path("static").exists():
    app.mount("/static", StaticFiles(directory="static"), name="static")


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("app:app", host="0.0.0.0", port=port, reload=True)
