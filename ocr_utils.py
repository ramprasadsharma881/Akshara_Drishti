"""
AksharaDrishti OCR Utilities - Cloud Vision Level Accuracy
===========================================================
Advanced OCR for old Telugu/Kannada Vedantha prints

Features:
1. Advanced preprocessing with OpenCV (CLAHE, denoising, adaptive threshold)
2. Line-by-line context padding for better Tesseract recognition
3. Confidence scoring with markers for low-confidence words
4. Specialized detection for Vedantic text patterns
5. Linguistic Repair Layer (Sandhi, Danda normalization, word joining)
"""

import cv2
import numpy as np
from PIL import Image
import pytesseract
from pytesseract import Output
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re
import os

# Limit OpenCV threads to prevent "Can't spawn new thread" errors
cv2.setNumThreads(2)
os.environ["OMP_NUM_THREADS"] = "2"


# ============================================================================
# CONFIGURATION
# ============================================================================

# Confidence threshold - words below this get marked with *
CONFIDENCE_THRESHOLD = 65

# Padding for better OCR recognition
LINE_PADDING = 30


# ============================================================================
# ADVANCED IMAGE PREPROCESSING - "CLEAN & POP" PIPELINE
# ============================================================================

def pil_to_cv2(pil_image):
    """Convert PIL Image to OpenCV format."""
    if pil_image.mode == 'RGB':
        return cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
    elif pil_image.mode == 'RGBA':
        return cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGBA2BGR)
    elif pil_image.mode == 'L':
        return np.array(pil_image)
    else:
        return cv2.cvtColor(np.array(pil_image.convert('RGB')), cv2.COLOR_RGB2BGR)


def cv2_to_pil(cv2_image):
    """Convert OpenCV image to PIL format."""
    if len(cv2_image.shape) == 2:  # Grayscale
        return Image.fromarray(cv2_image)
    else:  # BGR
        return Image.fromarray(cv2.cvtColor(cv2_image, cv2.COLOR_BGR2RGB))


def denoise_image(img_gray):
    """
    Apply Non-local Means Denoising.
    Removes noise from old paper without destroying thin curves.
    """
    return cv2.fastNlMeansDenoising(
        img_gray,
        None,
        h=10,
        templateWindowSize=7,
        searchWindowSize=21
    )


def apply_clahe(img_gray):
    """
    Apply CLAHE (Contrast Limited Adaptive Histogram Equalization).
    Critical for reading text near book binding where it gets dark.
    """
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    return clahe.apply(img_gray)


def adaptive_binarize(img_gray):
    """
    Apply Adaptive Gaussian Thresholding.
    Creates razor-sharp binary image for Tesseract's LSTM engine.
    """
    return cv2.adaptiveThreshold(
        img_gray,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        blockSize=15,
        C=8
    )


def add_context_padding(img_gray, padding=LINE_PADDING):
    """
    Add white padding around the image.
    Helps with Vattulu/subscripts recognition.
    """
    height, width = img_gray.shape
    padded = np.ones((height + 2 * padding, width + 2 * padding), dtype=np.uint8) * 255
    padded[padding:padding + height, padding:padding + width] = img_gray
    return padded


def preprocess_image_advanced(img_pil):
    """
    Advanced "Clean & Pop" preprocessing pipeline.
    """
    img = pil_to_cv2(img_pil)
    
    if len(img.shape) == 3:
        img_gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        img_gray = img
    
    # Pipeline: Denoise → CLAHE → Adaptive Threshold → Padding
    denoised = denoise_image(img_gray)
    clahe_applied = apply_clahe(denoised)
    binary = adaptive_binarize(clahe_applied)
    padded = add_context_padding(binary, padding=LINE_PADDING)
    
    return padded


def preprocess_image(img_pil):
    """
    Main preprocessing function.
    Crops headers/footers and applies advanced preprocessing.
    
    Header/footer cropping is aggressive to remove:
    - Book titles that appear on every page
    - Page numbers
    - Running headers
    """
    width, height = img_pil.size
    
    # More aggressive cropping to remove headers/footers
    # Headers often contain book title (repeated on every page)
    top_crop = int(height * 0.08)    # 8% from top (was 5%)
    bottom_crop = int(height * 0.06)  # 6% from bottom (was 4%)
    
    if height > 200:
        img_pil = img_pil.crop((0, top_crop, width, height - bottom_crop))
    
    processed = preprocess_image_advanced(img_pil)
    return processed


def preprocess_image_alternative(img_pil):
    """Alternative preprocessing - kept for compatibility."""
    return preprocess_image(img_pil)


# ============================================================================
# LINGUISTIC REPAIR LAYER
# ============================================================================

def normalize_dandas(text):
    """
    Danda Normalization - Fix OCR errors where । is read as I, l, or 1.
    
    Common OCR mistakes:
    - 'I' at end of sentence → '।'
    - 'l' at end of sentence → '।'
    - '1' at end of sentence → '।'
    - 'II' or '11' or 'll' → '॥'
    """
    if not text:
        return text
    
    # Double danda fixes (must be done before single danda)
    # Pattern: II, 11, ll, Il, lI, 1l, l1 at word boundaries → ॥
    text = re.sub(r'(?<=[।॥\s\u0C00-\u0CFF\u0900-\u097F])\s*[Il1]{2}\s*(?=[\s\d०-९೧-೯౧-౯]|$)', ' ॥ ', text)
    text = re.sub(r'\s+[Il1]{2}\s*$', ' ॥', text, flags=re.MULTILINE)
    
    # Single danda fixes
    # Pattern: I, l, or 1 at end of Indic word → ।
    text = re.sub(r'(?<=[\u0C00-\u0CFF\u0900-\u097F])\s*[Il1]\s*(?=\s|$)', '।', text)
    
    # Fix pipe characters that should be dandas
    text = re.sub(r'(?<=[\u0C00-\u0CFF\u0900-\u097F])\s*\|\s*(?=\s|$)', '।', text)
    text = re.sub(r'\|\|', '॥', text)
    
    # Clean up multiple dandas
    text = text.replace('।।', '॥')
    text = text.replace('। ।', '॥')
    
    # Normalize spacing around dandas
    text = re.sub(r'\s+।', '।', text)
    text = re.sub(r'\s+॥', '॥', text)
    text = re.sub(r'।\s+', '। ', text)
    text = re.sub(r'॥\s+', '॥ ', text)
    
    return text


def join_broken_words(text):
    """
    Sandhi and Word-Joiner Logic.
    
    Handles:
    1. Hyphenated words at line breaks: "Veda-\nnta" → "Vedanta"
    2. Orphaned vowel signs at start of new line
    3. Implicit Sandhi breaks
    """
    if not text:
        return text
    
    # Join hyphenated words at line breaks
    # Pattern: word ending with hyphen, followed by newline and continuation
    text = re.sub(r'-\s*\n\s*', '', text)
    
    # Join soft hyphens
    text = re.sub(r'­\s*\n\s*', '', text)  # Soft hyphen (U+00AD)
    
    # Telugu/Kannada vowel signs that might be orphaned at line start
    # These are dependent vowels (matras) that should attach to previous consonant
    vowel_signs_kannada = r'[ಾಿೀುೂೃೄೆೇೈೊೋೌ]'
    vowel_signs_telugu = r'[ాిీుూృౄెేైొోౌ]'
    vowel_signs_devanagari = r'[ािीुूृॄेैोौ]'
    
    # Join orphaned vowel signs to previous line
    pattern = f'\\s*\\n\\s*({vowel_signs_kannada}|{vowel_signs_telugu}|{vowel_signs_devanagari})'
    text = re.sub(pattern, r'\1', text)
    
    # Join consonant + virama at line end with next line consonant
    # Virama/Halant indicates the consonant should join with next
    virama_pattern = r'([\u0C4D\u094D\u0CCD])\s*\n\s*'  # Telugu, Devanagari, Kannada virama
    text = re.sub(virama_pattern, r'\1', text)
    
    return text


def remove_asterisk_artifacts(text):
    """
    Remove asterisks that are OCR artifacts.
    Tesseract sometimes misreads characters as *.
    
    Examples of artifacts:
    - *word* → word
    - *.0* → .0  
    - *ದಕ್ಷಿಣಾಮೂರ್ತಿ* → ದಕ್ಷಿಣಾಮೂರ್ತಿ
    """
    if not text:
        return text
    
    # Remove asterisks that wrap words: *word* → word
    text = re.sub(r'\*([^\s*]+)\*', r'\1', text)
    
    # Remove standalone asterisks surrounded by spaces
    text = re.sub(r'\s+\*\s+', ' ', text)
    
    # Remove asterisks at word boundaries
    text = re.sub(r'\*(?=[\u0C00-\u0CFF\u0900-\u097F\w])', '', text)  # *word → word
    text = re.sub(r'(?<=[\u0C00-\u0CFF\u0900-\u097F\w])\*', '', text)  # word* → word
    
    # Remove remaining isolated asterisks
    text = re.sub(r'(?<![*])\*(?![*])', '', text)
    
    return text


def clean_ocr_artifacts(text):
    """
    Remove common OCR artifacts and garbage characters.
    """
    if not text:
        return text
    
    # FIRST: Remove asterisk artifacts (very common OCR error)
    text = remove_asterisk_artifacts(text)
    
    # Remove isolated special characters that are likely noise
    text = re.sub(r'(?<!\S)[^\w\s\u0C00-\u0CFF\u0900-\u097F।॥\.,;:!?\-\'"()]+(?!\S)', '', text)
    
    # Remove repeated punctuation (except dandas)
    text = re.sub(r'[.,;:]{2,}', '.', text)
    
    # Clean up excessive whitespace
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove lines that are mostly garbage (< 50% valid characters)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            cleaned_lines.append('')
            continue
        
        valid_chars = len(re.findall(r'[\w\u0C00-\u0CFF\u0900-\u097F\s।॥,.]', line))
        if len(line) > 0 and (valid_chars / len(line) > 0.4 or len(line) < 5):
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)


def clean_ocr_text(text):
    """
    Complete text cleaning with Linguistic Repair Layer.
    """
    if not text:
        return ""
    
    # Step 1: Join broken words (Sandhi repair)
    text = join_broken_words(text)
    
    # Step 2: Normalize dandas (fix I/l/1 → ।)
    text = normalize_dandas(text)
    
    # Step 3: Clean OCR artifacts
    text = clean_ocr_artifacts(text)
    
    # Step 4: Final cleanup
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


# ============================================================================
# OCR WITH CONFIDENCE SCORING
# ============================================================================

def perform_ocr_with_confidence(img, language):
    """
    Perform OCR and return clean text with confidence statistics.
    NO asterisk markers - just clean text for Word document.
    
    Returns:
        tuple: (clean_text, low_confidence_count, avg_confidence)
    """
    if isinstance(img, np.ndarray):
        img_pil = Image.fromarray(img)
    else:
        img_pil = img
    
    config = '--oem 1 --psm 6 -c preserve_interword_spaces=1'
    
    try:
        data = pytesseract.image_to_data(
            img_pil, 
            lang=language, 
            config=config,
            output_type=Output.DICT
        )
        
        words = []
        confidences = []
        low_confidence_count = 0
        
        n_boxes = len(data['text'])
        current_line = -1
        line_words = []
        
        for i in range(n_boxes):
            word = data['text'][i].strip()
            if not word:
                continue
            
            conf = int(data['conf'][i]) if data['conf'][i] != '-1' else 0
            line_num = data['line_num'][i]
            
            if line_num != current_line:
                if line_words:
                    words.append(' '.join(line_words))
                    line_words = []
                current_line = line_num
            
            # Count low-confidence words (for statistics only)
            if conf < CONFIDENCE_THRESHOLD and conf > 0:
                low_confidence_count += 1
            
            # Add word WITHOUT asterisks - clean text only
            line_words.append(word)
            if conf > 0:
                confidences.append(conf)
        
        if line_words:
            words.append(' '.join(line_words))
        
        clean_text = '\n'.join(words)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        
        return clean_text, low_confidence_count, avg_confidence
        
    except Exception as e:
        text = pytesseract.image_to_string(img_pil, lang=language, config=config)
        return text, 0, 50


def perform_ocr(img, language):
    """Simple OCR wrapper for compatibility."""
    processed_text, _, _ = perform_ocr_with_confidence(img, language)
    return processed_text


def perform_ensemble_ocr(img, language):
    """Wrapper for compatibility."""
    return perform_ocr(img, language)


# ============================================================================
# TEXT CLASSIFICATION FOR FORMATTING
# ============================================================================

def is_shloka_or_verse(text):
    """
    Detect if text is a Shloka or Sanskrit verse.
    VERY CONSERVATIVE detection - only unmistakable verse patterns.
    
    A verse/shloka should be:
    1. A verse number line like ॥ 12 ॥
    2. A SHORT line ending with ॥ (verses are typically short, under 150 chars)
    3. Text starting with Om (ॐ)
    
    Regular paragraphs ending with ॥ are NOT verses - they are just sentences.
    """
    if not text or len(text) < 5:
        return False
    
    text = text.strip()
    
    # VERSE NUMBER PATTERNS - strongest indicator
    # Pattern: ॥ 12 ॥ or ॥ ೧೨ ॥ or just numbers between dandas
    if re.search(r'॥\s*[\d०-९೧-೯౧-౯]+\s*॥', text):
        # If it's JUST a verse number (very short), it's a verse marker
        if len(text) < 30:
            return True
    
    # Starts with Om - typically a mantra/shloka
    if re.match(r'^\s*(ॐ|ಓಂ|ఓం)\s*', text):
        return True
    
    # Short line ending with double danda - likely a verse
    # Long paragraphs ending with ॥ are just regular text, NOT verses
    if len(text) < 120 and re.search(r'॥\s*$', text):
        # Additional check: should not have too many words (verses are compact)
        word_count = len(text.split())
        if word_count < 15:
            return True
    
    # Text that starts AND ends with ॥ (enclosed verse)
    if re.match(r'^\s*॥', text) and re.search(r'॥\s*$', text):
        if len(text) < 150:
            return True
    
    # REMOVED: text ending with || - too ambiguous
    # REMOVED: any text ending with ॥ - regular sentences end with ॥ too
    
    return False


def get_verse_number(text):
    """
    Extract verse number from shloka if present.
    Returns the verse number string or None.
    """
    # Pattern: ॥ 12 ॥ or ॥ ೧೨ ॥ or ॥ १२ ॥
    match = re.search(r'॥\s*([\d०-९೧-೯౧-౯]+)\s*॥', text)
    if match:
        return match.group(1)
    return None


def is_uvacha_pattern(text):
    """
    Detect 'Uvacha' patterns (speaker introductions in Vedantic texts).
    """
    if not text or len(text) < 5:
        return False
    
    text = text.strip()
    
    uvacha_patterns = [
        r'उवाच\s*[।॥]?$',
        r'ಉವಾಚ\s*[।॥]?$',
        r'ఉవాచ\s*[।॥]?$',
        r'[A-Za-z]+\s+[Uu]vacha',
        r'श्री\s*\S+\s*उवाच',
        r'ಶ್ರೀ\s*\S+\s*ಉವಾಚ',
        r'శ్రీ\s*\S+\s*ఉవాచ',
        r'भगवान्?\s*उवाच',
        r'अर्जुन\s*उवाच',
        r'संजय\s*उवाच',
    ]
    
    for pattern in uvacha_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    
    return False


def is_heading_or_title(text):
    """
    Detect if text is a heading or title.
    VERY CONSERVATIVE detection - only unmistakable chapter/section headers.
    
    Headings should be:
    1. Very short (typically under 50 chars)
    2. Contain explicit chapter/section keywords
    3. Not contain regular sentence punctuation
    """
    if not text or len(text) < 3:
        return False
    
    text = text.strip()
    
    # Long text is NEVER a heading
    if len(text) > 80:
        return False
    
    # Text with sentence-ending punctuation is NOT a heading
    # (headings don't have periods, commas in the middle, etc.)
    if re.search(r'[।॥,;]\s+\S', text):  # Punctuation followed by more text = sentence, not heading
        return False
    
    # Shlokas are not headings
    if is_shloka_or_verse(text):
        return False
    
    # Uvacha patterns (speaker introductions) - these are special
    if is_uvacha_pattern(text):
        return True
    
    # Chapter/Section keywords in various scripts
    chapter_patterns = [
        r'अध्याय|प्रकरण|खण्ड|भाग|परिच्छेद',  # Sanskrit/Hindi
        r'ಅಧ್ಯಾಯ|ಪ್ರಕರಣ|ಖಂಡ|ಭಾಗ|ಪರಿಚ್ಛೇದ',  # Kannada
        r'అధ్యాయ|ప్రకరణ|ఖండ|భాగ|పరిచ్ఛేద',  # Telugu
    ]
    for pattern in chapter_patterns:
        if re.search(pattern, text):
            # Extra check: should be short
            if len(text) < 60:
                return True
    
    # Numbered chapter at start: "1.", "೧." etc - ONLY if very short
    if re.match(r'^\s*[\d०-९೧-೯౧-౯]+[\.\)]\s*\S', text):
        if len(text) < 50:  # Very short numbered items only
            return True
    
    return False


def is_sanskrit_content(text):
    """
    Detect if text is PRIMARILY Sanskrit (Devanagari) content.
    This should be VERY conservative - only pure Sanskrit shlokas should be bolded.
    
    Regular commentary in Kannada/Telugu that quotes Sanskrit should NOT be all bold.
    """
    if not text or len(text) < 10:
        return False
    
    # Only consider this for SHORT text (likely standalone Sanskrit verses)
    # Long paragraphs are usually commentary, not pure Sanskrit
    if len(text) > 150:
        return False
    
    devanagari = len(re.findall(r'[\u0900-\u097F]', text))
    kannada = len(re.findall(r'[\u0C80-\u0CFF]', text))
    telugu = len(re.findall(r'[\u0C00-\u0C7F]', text))
    total_text = text.replace(' ', '').replace('\n', '')
    
    if len(total_text) > 0:
        # Only bold if it's PREDOMINANTLY Devanagari (>70%)
        # AND not mixed with significant Kannada/Telugu
        ratio = devanagari / len(total_text)
        local_script_ratio = (kannada + telugu) / len(total_text)
        
        # Must be mostly Sanskrit AND not mixed with local script
        if ratio > 0.7 and local_script_ratio < 0.2:
            return True
    
    return False


# ============================================================================
# SINGLE PAGE PROCESSING FUNCTION (For Parallel Processing)
# ============================================================================

def process_single_page(args):
    """
    Process a single page - designed for parallel execution.
    Includes robust error handling so one failed page doesn't crash the book.
    
    Args:
        args: tuple of (page_idx, img_pil, language)
    
    Returns:
        dict with page_idx, paragraphs, low_confidence_count, success status
    """
    page_idx, img_pil, language = args
    
    try:
        # Preprocess the image
        processed = preprocess_image(img_pil)
        
        # Perform OCR with confidence scoring
        raw_text, low_conf_count, confidence = perform_ocr_with_confidence(processed, language)
        
        # Apply Linguistic Repair Layer
        cleaned_text = clean_ocr_text(raw_text)
        
        if not cleaned_text.strip():
            return {
                'page_idx': page_idx,
                'paragraphs': [],
                'low_confidence_count': 0,
                'confidence': 0,
                'success': True
            }
        
        # Process paragraphs
        paragraphs = []
        raw_paragraphs = cleaned_text.split('\n\n')
        
        for p_text in raw_paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue
            
            lines = p_text.split('\n')
            
            # Check if verse-like structure
            is_verse_structure = False
            non_empty_lines = [l.strip() for l in lines if l.strip()]
            if non_empty_lines:
                danda_endings = sum(1 for l in non_empty_lines if re.search(r'[।॥]$', l))
                if danda_endings >= len(non_empty_lines) * 0.5 and len(non_empty_lines) >= 2:
                    is_verse_structure = True
            
            if is_verse_structure:
                clean_text = '\n'.join(l.strip() for l in lines if l.strip())
            else:
                clean_text = ' '.join(l.strip() for l in lines if l.strip())
            
            if not clean_text or len(clean_text) < 2:
                continue
            
            # Classify text
            para_data = {
                'text': clean_text,
                'is_shloka': is_shloka_or_verse(clean_text),
                'is_heading': is_heading_or_title(clean_text),
                'is_uvacha': is_uvacha_pattern(clean_text),
                'is_sanskrit': is_sanskrit_content(clean_text),
                'verse_number': get_verse_number(clean_text)
            }
            paragraphs.append(para_data)
        
        return {
            'page_idx': page_idx,
            'paragraphs': paragraphs,
            'low_confidence_count': low_conf_count,
            'confidence': confidence,
            'success': True
        }
        
    except Exception as e:
        # Error handling - return failure info but don't crash
        return {
            'page_idx': page_idx,
            'paragraphs': [],
            'low_confidence_count': 0,
            'confidence': 0,
            'success': False,
            'error': str(e)
        }


# ============================================================================
# DOCUMENT BUILDING FUNCTION
# ============================================================================

def add_paragraphs_to_document(doc, paragraphs):
    """
    Add processed paragraphs to the document with proper formatting.
    
    Formatting rules:
    1. All text is LEFT JUSTIFIED
    2. Headings and subheadings are BOLD
    3. Shlokas and Sanskrit sentences are BOLD in separate paragraphs
    4. Uvacha patterns are BOLD
    5. Minimal paragraph spacing
    6. Verse numbers get extra spacing
    """
    for para_data in paragraphs:
        text = para_data['text']
        is_shloka = para_data['is_shloka']
        is_heading = para_data['is_heading']
        is_uvacha = para_data['is_uvacha']
        is_sanskrit = para_data['is_sanskrit']
        verse_number = para_data.get('verse_number')
        
        # Create paragraph
        paragraph = doc.add_paragraph()
        
        # LEFT JUSTIFY (Requirement 1)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Minimal spacing (Requirement 4)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15
        
        # Add text
        run = paragraph.add_run(text)
        
        # Apply formatting based on content type
        if is_heading or is_uvacha:
            run.bold = True
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)
        elif is_shloka:
            run.bold = True
            paragraph_format.space_before = Pt(8)
            paragraph_format.space_after = Pt(8)
            # Extra spacing for verses with numbers
            if verse_number:
                paragraph_format.space_before = Pt(10)
                paragraph_format.space_after = Pt(10)
        elif is_sanskrit:
            run.bold = True


def add_quality_report(doc, total_low_confidence, total_pages, avg_confidence, failed_pages):
    """
    Add OCR Quality Report at the end of the document.
    """
    # Add page break before report
    doc.add_page_break()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("═" * 50)
    title_run.bold = True
    
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2_run = title2.add_run("OCR QUALITY REPORT")
    title2_run.bold = True
    title2_run.font.size = Pt(16)
    
    title3 = doc.add_paragraph()
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title3_run = title3.add_run("═" * 50)
    title3_run.bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Statistics
    stats = [
        f"Total Pages Processed: {total_pages}",
        f"Average OCR Confidence: {avg_confidence:.1f}%",
        f"",
        f"Low-Confidence Words Found: {total_low_confidence}",
        f"(These words may need manual verification)",
        f"",
    ]
    
    for stat in stats:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(stat)
        if "Low-Confidence" in stat:
            run.bold = True
    
    # Failed pages (if any)
    if failed_pages:
        doc.add_paragraph()
        fail_title = doc.add_paragraph()
        fail_run = fail_title.add_run("Pages That Failed to Process:")
        fail_run.bold = True
        
        for page_num, error in failed_pages:
            p = doc.add_paragraph()
            p.add_run(f"  • Page {page_num + 1}: {error}")
    
    # Footer note
    doc.add_paragraph()
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("─" * 40)
    
    note2 = doc.add_paragraph()
    note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note2_run = note2.add_run("Generated by AksharaDrishti OCR System")
    note2_run.italic = True
    
    note3 = doc.add_paragraph()
    note3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note3_run = note3.add_run("For Vedantic Literature Preservation")
    note3_run.italic = True


def process_image_to_docx_content(doc, img, language, config):
    """
    Process image and add formatted content to document.
    For backward compatibility with sequential processing.
    """
    if isinstance(img, np.ndarray):
        img_pil = Image.fromarray(img)
    else:
        img_pil = img
    
    result = process_single_page((0, img_pil, language))
    
    if result['success'] and result['paragraphs']:
        add_paragraphs_to_document(doc, result['paragraphs'])
